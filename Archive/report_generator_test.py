import pandas as pd

def read_and_process_data(csv_file):
    # Read the CSV file
    df = pd.read_csv(csv_file)
    
    # Perform simple data analysis
    # Example: Calculate average of a column
    df['average'] = df.mean(axis=1)
    
    return df

import matplotlib.pyplot as plt

def create_charts(df, chart_path):
    # Example: Create a bar chart
    plt.figure(figsize=(10, 6))
    df['average'].plot(kind='bar')
    plt.title('Average Values')
    plt.xlabel('Index')
    plt.ylabel('Average')
    plt.savefig(chart_path)
    plt.close()

def create_tables(df):
    # Convert DataFrame to table (list of lists)
    table = df.values.tolist()
    header = df.columns.tolist()
    return header, table

from docx import Document
from docx.shared import Inches

def create_word_report(title, df, chart_path, output_path):
    doc = Document()
    
    # Title Page
    doc.add_heading(title, 0)
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Data Analysis\n2. Charts and Tables\n3. References')
    doc.add_page_break()
    
    # Data Analysis Section
    doc.add_heading('Data Analysis', level=1)
    doc.add_paragraph('This section contains data analysis.')
    
    # Charts Section
    doc.add_heading('Charts and Tables', level=1)
    doc.add_paragraph('This section contains charts and tables.')
    doc.add_picture(chart_path, width=Inches(6))
    
    # Insert table
    header, table = create_tables(df)
    table = doc.add_table(rows=1, cols=len(header))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(header):
        hdr_cells[i].text = col
    for row in table:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    
    # References Section
    doc.add_heading('References', level=1)
    doc.add_paragraph('References go here.')
    
    # Save the document
    doc.save(output_path)
    
def generate_report(csv_file, title, output_dir):
    # Process data
    df = read_and_process_data(csv_file)
    
    # Create chart
    chart_path = output_dir + '/chart.png'
    create_charts(df, chart_path)
    
    # Create Word report
    output_path = output_dir + '/report.docx'
    create_word_report(title, df, chart_path, output_path)
