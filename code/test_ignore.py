import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# Load the Excel workbook
input_path = 'test_input.xlsx'  # Update with the actual path to your Excel file

input_workbook = pd.read_excel(input_path, engine = 'openpyxl', sheet_name=None)

utility_bills = pd.DataFrame(input_workbook['Utility Bills'])
# Load data from Excel sheets
df_electricity = utility_bills.filter(regex='kw|peak')
df_water = utility_bills.filter(regex='Water|Sewer')

# Create a Word document
doc = Document()

# Add Title
doc.add_heading('Energy and Utility Bill Analysis', level=1)

# Add Introduction
doc.add_paragraph(
    "Utility-bill analysis for the facility was conducted by compiling 12 months of electricity bills from "
    "North Georgia Electric Membership Corp (Account# 447888001) and 12 months of water bills from Dalton "
    "Utilities (Account#: 0005668601)."
)

# Add Table D.1: Monthly Electricity and Water Consumptions and $ Charges
doc.add_heading('Table D.1: Monthly Electricity and Water Consumptions and $ Charges', level=2)

table1 = doc.add_table(rows=1, cols=len(df_electricity.columns))
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Add header row.
hdr_cells = table1.rows[0].cells
for i, column_name in enumerate(df_electricity.columns):
    hdr_cells[i].text = column_name

# Add the rest of the data rows.
for index, row in df_electricity.iterrows():
    row_cells = table1.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Add space between tables
doc.add_paragraph()

# Add Table D.2: Monthly Electricity Rate Charges
doc.add_heading('Table D.2: Monthly Electricity Rate Charges', level=2)

table2 = doc.add_table(rows=1, cols=len(df_water.columns))
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Add header row.
hdr_cells = table2.rows[0].cells
for i, column_name in enumerate(df_water.columns):
    hdr_cells[i].text = column_name

# Add the rest of the data rows.
for index, row in df_water.iterrows():
    row_cells = table2.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Add space before analysis paragraph
doc.add_paragraph()

# Add Analysis Paragraph
doc.add_paragraph(
    "The analysis of electricity bills shows that the facility is charged based on less or above 15000 kWh usage. "
    "Table D.2 outlines the detailed breakdown of electricity charges for the year 2023 incorporating energy charges "
    "TVA (Tennessee Valley Authority) Power Generation Fuel Charges and demand charges. The first two columns "
    "distinguish between energy charges for usage below and above 15000 kilowatt-hours (kWh) respectively. The subsequent "
    "columns provide insights into TVA Power Generation Fuel Charges with distinctions for usage ranges. Additionally, demand "
    "charges are specified for both consumption categories.\n\n"
    "In January to March, the energy charges remain consistent at $ 0.08226 for usage below 15000 kWh and $ 0.03511 for usage above "
    "15000 kWh. The TVA Power Generation Fuel Charges exhibit a declining trend starting at $ 0.0349 for both usage categories in "
    "January and February, dropping to $ 0.02741 and $ 0.02182 in March. Subsequent months follow a similar pattern with May seeing "
    "a slight increase in TVA Power Generation Fuel Charges. Notably, June introduces an increase in both energy charges and TVA Power "
    "Generation Fuel Charges. The demand charges persist at $ 14.86 until May, after which they rise to $ 15.81, maintaining consistency "
    "throughout the rest of the year. The latter part of the year sees marginal variations in energy charges and TVA Power Generation "
    "Fuel Charges, with demand charges experiencing slight increases in November and December. Overall, the table illustrates a nuanced "
    "interplay of energy pricing components reflecting adjustments influenced by factors such as fuel costs and energy demand.\n\n"
    "For this report, we have used the blended $0.086 per kWh + $15.44 per Peak KW rate for electricity reduction improvements in Section E. "
    "The blended $ per kWh is the total annual electricity $ charges divided by total annual kWh consumption. For the water charges, an average "
    "rate of $5.65 per thousand gallons of water was found based on the Dalton Utilities bills."
)

# Save the document
output_path = 'test_report.docx'
doc.save(output_path)

print(f"Report generated and saved as {output_path}")
